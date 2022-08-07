from docx import Document


def create_invoice(info):
	doc = Document()
	doc.add_heading(f"Invoice from {info['client_name']} to {info['vendor_name']}")

	name = doc.add_paragraph(f"Client name: ")
	name.add_run(info['client_name']).bold = True

	vendor = doc.add_paragraph(f"Vendor name: ")
	vendor.add_run(info['vendor_name']).bold = True

	amount = doc.add_paragraph(f"Amount: ")
	amount.add_run(info["amount"]+ " DA").bold = True

	quantity = doc.add_paragraph(f"Quantity: ")
	quantity.add_run(info["quantity"]).bold = True

	doc.save("invoice.docx")